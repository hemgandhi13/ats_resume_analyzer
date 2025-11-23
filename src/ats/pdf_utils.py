from __future__ import annotations

import io
import logging
from typing import TYPE_CHECKING, Callable, List, Optional, Union

import fitz  # PyMuPDF

try:
    import docx  # python-docx
except Exception:
    docx = None  # type: ignore

if TYPE_CHECKING:  # imports only for type checkers, not at runtime
    # Pillow's Image type for annotations
    from PIL.Image import Image as PILImage  # type: ignore

try:
    from PIL import Image  # type: ignore
except Exception:  # Pillow not installed
    Image = None  # type: ignore

try:
    import pytesseract  # type: ignore
except Exception:  # pytesseract not installed
    pytesseract = None  # type: ignore

logger = logging.getLogger(__name__)


# ============================================================
# Small text helpers
# ============================================================


def normalize(s: str) -> str:
    """Lowercase + collapse whitespace"""
    import re as _re

    return _re.sub(r"\s+", " ", s.lower()).strip()


def tokenize_words(s: str) -> List[str]:
    """Split into alphanumeric-ish tokens; good enough for skills detection."""
    import re as _re

    return _re.findall(r"[a-zA-Z0-9+#\-]+", s.lower())


def _render_page_to_image(page: "fitz.Page", dpi: int = 300) -> "PILImage":
    """
    Render a PDF page to a Pillow Image at the given DPI.
    """
    if Image is None:
        raise RuntimeError("Pillow is not installed. Add 'pillow' to dependencies.")
    scale = dpi / 72.0
    mat = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=mat, alpha=False)  # RGB
    img_bytes = pix.tobytes("png")
    return Image.open(io.BytesIO(img_bytes))


def _ocr_image(img: "PILImage", lang: str = "eng") -> str:
    """
    Run Tesseract OCR on a Pillow Image.
    """
    if pytesseract is None:
        raise RuntimeError("pytesseract is not installed. Add 'pytesseract' to dependencies.")
    return pytesseract.image_to_string(img, lang=lang)


def extract_text_from_pdf(
    data: Union[str, bytes, bytearray],
    *,
    password: Optional[str] = None,
    max_pages: Optional[int] = None,
    return_pages: bool = False,
    enable_ocr: bool = True,
    ocr_lang: str = "eng",
    ocr_dpi: int = 300,
    ocr_page_char_threshold: int = 40,
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
    tesseract_cmd: Optional[str] = None,
) -> Union[str, List[str]]:
    """
    Extract text from a PDF with optional OCR fallback.

    Args:
        data: File path (str) or raw PDF bytes.
        password: Password for encrypted PDFs (if required).
        max_pages: Stop after this many pages (from the beginning).
        return_pages: If True, return a list of per-page texts; else a single joined string.
        enable_ocr: Try OCR when extracted text on a page looks too short.
        ocr_lang: Tesseract language code(s), e.g., "eng" or "eng+deu".
        ocr_dpi: Rasterization DPI for OCR rendering (300 is a good default).
        ocr_page_char_threshold: If page text length < threshold, attempt OCR.
        progress_callback: Optional fn(page_index, total_pages, message) for UI progress.
        tesseract_cmd: Optional absolute path to tesseract executable (Windows).

    Returns:
        str or List[str]: Extracted text (joined or per-page).
    """
    # Configure pytesseract path on Windows if provided
    if tesseract_cmd and pytesseract is not None:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd  # type: ignore

    # Open the document from path or bytes
    if isinstance(data, (bytes, bytearray)):
        doc = fitz.open(stream=data, filetype="pdf")
    elif isinstance(data, str):
        doc = fitz.open(data)
    else:
        raise TypeError("data must be a file path (str) or PDF bytes/bytearray")

    try:
        # Handle encrypted PDFs
        if doc.needs_pass:
            if not password or not doc.authenticate(password):
                raise ValueError("PDF is encrypted and no/invalid password was provided.")

        total_pages = doc.page_count
        limit = min(total_pages, max_pages) if max_pages else total_pages

        results: List[str] = []
        for i in range(limit):
            if progress_callback:
                progress_callback(i, limit, "extracting page text")

            page = doc.load_page(i)
            text = page.get_text("text") or ""

            # If the page seems image-only (very short text), attempt OCR
            if enable_ocr and len(text.strip()) < ocr_page_char_threshold:
                try:
                    if progress_callback:
                        progress_callback(i, limit, "running OCR")
                    img = _render_page_to_image(page, dpi=ocr_dpi)
                    ocr_text = _ocr_image(img, lang=ocr_lang)
                    if ocr_text and len(ocr_text.strip()) > len(text.strip()):
                        text = ocr_text
                except Exception as e:
                    logger.warning("OCR failed on page %s: %s", i, e)

            results.append(text)

        if progress_callback:
            progress_callback(limit, limit, "done")

        return results if return_pages else "\n".join(results)
    finally:
        doc.close()


def extract_text_from_docx(data: Union[str, bytes, bytearray]) -> str:
    """
    Read DOCX from path or bytes and return plain text.
    """
    if docx is None:
        raise RuntimeError("python-docx is not installed. Add 'python-docx' to dependencies.")

    if isinstance(data, str):
        d = docx.Document(data)
    elif isinstance(data, (bytes, bytearray)):
        import io

        d = docx.Document(io.BytesIO(data))
    else:
        raise TypeError("data must be a file path (str) or DOCX bytes/bytearray")

    parts = []
    for p in d.paragraphs:
        if p.text:
            parts.append(p.text)
    # Also harvest simple table text
    for table in d.tables:
        for row in table.rows:
            row_txt = [cell.text for cell in row.cells if cell.text]
            if row_txt:
                parts.append(" | ".join(row_txt))

    return "\n".join(parts)


def extract_text_from_txt(data: Union[str, bytes, bytearray], encoding: str = "utf-8") -> str:
    """
    Read TXT/MD from path or bytes and return text.
    """
    if isinstance(data, str):
        with open(data, "r", encoding=encoding, errors="ignore") as f:
            return f.read()
    elif isinstance(data, (bytes, bytearray)):
        return data.decode(encoding, errors="ignore")
    else:
        raise TypeError("data must be a file path (str) or text bytes/bytearray")


def extract_text_auto(
    data: Union[str, bytes, bytearray],
    *,
    filename: Optional[str] = None,
    **kwargs,
) -> str:
    """
    Auto-detect based on file extension (from `filename`) and extract text.
    Supported: PDF (with OCR fallback via extract_text_from_pdf), DOCX, TXT/MD.
    Extra kwargs are forwarded to extract_text_from_pdf (e.g., enable_ocr, ocr_lang),
    but this helper always returns a single joined string (not per-page).
    """
    ext = ""
    if filename:
        ext = filename.rsplit(".", 1)[-1].lower()

    # For PDF calls, enforce return_pages=False so we always get a str
    pdf_kwargs = dict(kwargs)
    pdf_kwargs.pop("return_pages", None)

    if ext == "pdf" or (not ext and isinstance(data, (bytes, bytearray))):
        text = extract_text_from_pdf(data, return_pages=False, **pdf_kwargs)
        # At runtime this is guaranteed by return_pages=False; assert for type checkers
        assert isinstance(text, str)
        return text

    if ext == "docx":
        return extract_text_from_docx(data)

    if ext in {"txt", "md"}:
        return extract_text_from_txt(data)

    # Fallback: try PDF first, then text decode
    try:
        text = extract_text_from_pdf(data, return_pages=False, **pdf_kwargs)
        assert isinstance(text, str)
        return text
    except Exception:
        return extract_text_from_txt(data)
